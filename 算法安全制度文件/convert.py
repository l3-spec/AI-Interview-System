#!/usr/bin/env python3
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RULES_DIR = "/Users/linxiong/Documents/GitHub/AI-Interview-System/算法安全制度文件"

def md_to_text(md_content):
    lines = md_content.split('\n')
    result = []
    in_table = False
    table_lines = []

    for line in lines:
        line = line.rstrip()

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            table_lines.append(line)
            if re.match(r'^[\|: -]+$', line.replace('|', '').strip()):
                in_table = False
                table_lines = []
            continue
        else:
            if table_lines:
                for tl in table_lines:
                    result.append(tl)
                result.append('')
                table_lines = []

            # Headings
            m = re.match(r'^(#{1,6})\s+(.*)', line)
            if m:
                result.append(m.group(2))
                continue

            # Bold
            line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            # Italic
            line = re.sub(r'\*(.+?)\*', r'\1', line)
            # Code
            line = re.sub(r'`(.+?)`', r'\1', line)

            result.append(line)

    return '\n'.join(result)


def apply_heading_style(paragraph, level):
    style_name = {
        1: 'Title',
        2: 'Heading 1',
        3: 'Heading 2',
        4: 'Heading 3',
    }.get(level, 'Normal')
    if paragraph.style.name != style_name:
        paragraph.style = style_name


def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract front-matter info
    title = ''
    version = ''
    effective_date = ''

    lines = content.split('\n')
    body_lines = []
    i = 0
    for line in lines:
        m = re.match(r'\*\*(.+?)\*\*[：:]\s*(.+)', line)
        if m:
            key, val = m.group(1), m.group(2)
            if '文件编号' in key:
                version = val
            elif '生效日期' in key:
                effective_date = val
        if i < 5 and re.match(r'^#{1,3}\s', line):
            title = re.sub(r'^#{1,3}\s+', '', line).strip()
            i += 1
            continue
        body_lines.append(line)
        i += 1

    full_text = '\n'.join(body_lines)
    doc = Document()

    # Set default font for Chinese
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Normal'].font.size = Pt(11)

    # Title paragraph
    title_para = doc.add_heading('', level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title if title else os.path.basename(md_path))
    run.bold = True
    run.font.size = Pt(18)

    # Metadata block
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_lines = []
    for line in lines[:20]:
        m = re.match(r'\*\*(.+?)\*\*[：:]\s*(.+)', line)
        if m:
            meta_lines.append(f"{m.group(1)}：{m.group(2)}")
    if meta_lines:
        meta_para.add_run('\n'.join(meta_lines))

    doc.add_paragraph('')

    in_table = False
    table_rows = []

    for line in full_text.split('\n'):
        line = line.rstrip()

        # Table rows
        if '|' in line and line.strip().startswith('|') and not re.match(r'^[\|: -]+$', line.replace('|', '').strip()):
            # Skip separator rows and empty cells
            cells = [c.strip() for c in line.split('|')]
            if len(cells) > 1:
                table_rows.append(cells[1:-1] if line.strip().startswith('|') else cells)
            continue
        else:
            if table_rows:
                tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                tbl.style = 'Table Grid'
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = cell_text
                doc.add_paragraph('')
                table_rows = []

        # Skip metadata already shown
        if re.match(r'\*\*(文件编号|版本|生效日期|制定单位)\*\*', line):
            continue

        # Headings
        hm = re.match(r'^(#{1,6})\s+(.*)', line)
        if hm:
            level = len(hm.group(1))
            h = doc.add_heading(hm.group(2), level=min(level, 3))
            continue

        # Bold text
        parts = re.split(r'\*\*(.+?)\*\*', line)
        if len(parts) > 1:
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    run = p.add_run(part)
                else:
                    run = p.add_run(part)
                    run.bold = True
        elif line:
            p = doc.add_paragraph(line)
            p.style = doc.styles['Normal']

    doc.save(docx_path)


# File list in order
files = [
    ("01_算法安全管理办法.md", "01_算法安全管理办法.docx"),
    ("02_算法安全自评估制度.md", "02_算法安全自评估制度.docx"),
    ("03_算法安全监测制度.md", "03_算法安全监测制度.docx"),
    ("04_算法安全事件应急预案.md", "04_算法安全事件应急预案.docx"),
    ("05_算法违法违规处置制度.md", "05_算法违法违规处置制度.docx"),
    ("06_数据安全管理制度.md", "06_数据安全管理制度.docx"),
    ("07_用户个人信息保护制度.md", "07_用户个人信息保护制度.docx"),
    ("08_科技伦理审查制度.md", "08_科技伦理审查制度.docx"),
    ("09_算法研发安全管理制度.md", "09_算法研发安全管理制度.docx"),
    ("10_用户投诉处理制度.md", "10_用户投诉处理制度.docx"),
    ("11_AI面试系统安全技术方案.md", "11_AI面试系统安全技术方案.docx"),
    ("12_AI面试系统数据安全方案.md", "12_AI面试系统数据安全方案.docx"),
]

for md_name, docx_name in files:
    md_path = os.path.join(RULES_DIR, md_name)
    docx_path = os.path.join(RULES_DIR, docx_name)
    if os.path.exists(md_path):
        md_to_docx(md_path, docx_path)
        print(f"✓ {docx_name}")
    else:
        print(f"✗ {md_name} not found")

print("Done!")